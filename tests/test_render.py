"""Golden typography test plus structural rule tests for specs/03-renderer.md.

render.py does not exist yet. Every test here is expected to fail on
collection until it is implemented.

_EXPECTED below is hand-transcribed from two independent sources, neither of
which is render.py or a prior render() call:
  1. specs/03-renderer.md's own typography table.
  2. Raw XML pulled directly from docs/headless-headhunter/template.docx
     (byte-identical to resume/templates/golden.docx): rFonts Arial
     throughout, w:sz values 21/24/28 (10.5/12/14pt), pgMar 720 twips all
     four sides, explicit w:spacing line=276 auto (1.15) on some paragraphs
     and the document's own line=360 auto default (1.5) elsewhere, and a
     w:tab right@10800 (7.5in) already present among the template's three
     inconsistent tab configs.
The two values that deliberately differ from the raw template (left
alignment instead of w:jc="both", one tab stop instead of three) are spec
03's own named fixes, not something this table invents.
"""

import hashlib
from dataclasses import dataclass
from pathlib import Path

import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor, Twips

from jobengine.resume.bank import (
    Bank,
    Bullet,
    Education,
    Meta,
    Publication,
    Role,
    SummaryBullet,
    load_bank,
)
from jobengine.resume.render import Identity, RenderProfile, render

_TEMPLATE_PATH = Path("resume/templates/golden.docx")
_IDENTITY_PATH = Path("identity.toml")


@dataclass(frozen=True)
class Typography:
    font: str
    size: Pt
    bold: bool
    italic: bool
    alignment: WD_ALIGN_PARAGRAPH
    line_spacing: float


_EXPECTED = {
    "name": Typography("Arial", Pt(14), True, False, WD_ALIGN_PARAGRAPH.CENTER, 1.15),
    "contact": Typography(
        "Arial", Pt(12), False, False, WD_ALIGN_PARAGRAPH.CENTER, 1.15
    ),
    "status": Typography(
        "Arial", Pt(12), False, False, WD_ALIGN_PARAGRAPH.CENTER, 1.15
    ),
    "section_header": Typography(
        "Arial", Pt(10.5), True, False, WD_ALIGN_PARAGRAPH.LEFT, 1.5
    ),
    "job_title_and_date": Typography(
        "Arial", Pt(10.5), False, True, WD_ALIGN_PARAGRAPH.LEFT, 1.5
    ),
    "bullet": Typography(
        "Arial", Pt(10.5), False, False, WD_ALIGN_PARAGRAPH.LEFT, 1.5
    ),
    "publication": Typography(
        "Arial", Pt(10.5), False, False, WD_ALIGN_PARAGRAPH.LEFT, 1.5
    ),
}
_EXPECTED_MARGIN = Twips(720)
_EXPECTED_TAB_POSITION = Twips(10800)


def _assert_run_matches(run, expected: Typography) -> None:
    assert run.font.name == expected.font
    assert run.font.size == expected.size
    assert run.font.bold == expected.bold
    assert run.font.italic == expected.italic


def _assert_paragraph_matches(paragraph, expected: Typography) -> None:
    assert paragraph.alignment == expected.alignment
    assert paragraph.paragraph_format.line_spacing == pytest.approx(
        expected.line_spacing
    )
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        _assert_run_matches(run, expected)


def _find_paragraph(doc, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise AssertionError(f"no paragraph with text {text!r}")


def _paragraphs_after(doc, anchor_text: str, count: int):
    texts = [p.text.strip() for p in doc.paragraphs]
    idx = texts.index(anchor_text)
    return doc.paragraphs[idx + 1 : idx + 1 + count]


# ---------------------------------------------------------------------------
# Synthetic fixtures, small and fast. One real education entry, one
# full_time role (summary + 2 bullets, current, "to Present"), one project
# role (summary + 2 bullets, no dates), one publication. Enough to exercise
# every semantic role and every structural rule once.
# ---------------------------------------------------------------------------


def _synthetic_identity(**overrides) -> Identity:
    fields = {
        "full_name": "Test Person",
        "email": "test@example.com",
        "phone": "+1-555-000-0000",
        "city": "Austin",
        "state": "TX",
        "linkedin": "https://www.linkedin.com/in/testperson/",
        "github": "https://github.com/testperson",
        "portfolio": "https://testperson.dev/",
        "scholar": "https://scholar.google.com/citations?user=test",
        "work_authorization_statement": "Authorized to work in the US",
    }
    fields.update(overrides)
    return Identity(**fields)


def _synthetic_bank(*, education: list[Education] | None = None) -> Bank:
    if education is None:
        education = [
            Education(
                id="edu_test",
                degree="MS",
                field="Computer Engineering",
                institution="Test University",
                gpa="3.9/4.0",
                status="May 2025",
                requires_degree_profiles=["ai_ml_engineer"],
            )
        ]
    full_time_role = Role(
        id="role_test_ft",
        company="TestCo",
        location="Austin, TX",
        start="2026-03",
        end=None,
        kind="full_time",
        title={"default": "Test Engineer"},
        summary=SummaryBullet(
            id="b_test_ft_sum",
            text="Built the systems behind a test product",
            keywords=[],
            status="verified",
        ),
        bullets=[
            Bullet(
                id="b_test_ft_01",
                status="verified",
                what="a thing",
                how="a way",
                result="a result",
                text="Built a thing using a way that produced a result",
                keywords=[],
                evidence="internal",
                profiles=["ai_ml_engineer"],
            ),
            Bullet(
                id="b_test_ft_02",
                status="verified",
                what="another thing",
                how="another way",
                result="another result",
                text="Shipped another thing using another way that produced another result",
                keywords=[],
                evidence="internal",
                profiles=["ai_ml_engineer"],
            ),
        ],
    )
    project_role = Role(
        id="role_test_project",
        kind="project",
        title={"default": "Test Side Project"},
        summary=SummaryBullet(
            id="b_test_proj_sum",
            text="Built a small tool for personal use",
            keywords=[],
            status="verified",
        ),
        bullets=[
            Bullet(
                id="b_test_proj_01",
                status="verified",
                what="a feature",
                how="a technique",
                result="an outcome",
                text="Built a feature using a technique that produced an outcome",
                keywords=[],
                evidence="internal",
                profiles=["ai_ml_engineer"],
            ),
            Bullet(
                id="b_test_proj_02",
                status="verified",
                what="a second feature",
                how="a second technique",
                result="a second outcome",
                text="Added a second feature using a second technique that produced a second outcome",
                keywords=[],
                evidence="internal",
                profiles=["ai_ml_engineer"],
            ),
        ],
    )
    publication = Publication(
        id="pub_test",
        text="A test paper about testing things.",
        authors_bold="Person, T.",
        venue="Journal of Testing",
        url="https://example.com/paper",
    )
    return Bank(
        meta=Meta(owner="Test", updated="2026-01-01"),
        education=education,
        roles=[full_time_role, project_role],
        publications=[publication],
    )


def _default_profile(**overrides) -> RenderProfile:
    fields = {
        "section_order": ["education", "work_history", "projects", "publications"],
        "include_summary": False,
        "summary_text": None,
    }
    fields.update(overrides)
    return RenderProfile(**fields)


# ---------------------------------------------------------------------------
# Golden typography tests, per semantic role, against the synthetic bank.
# ---------------------------------------------------------------------------


def test_name_typography():
    identity = _synthetic_identity()
    doc = render(_synthetic_bank(), identity, _default_profile())
    paragraph = _find_paragraph(doc, identity.full_name)
    _assert_paragraph_matches(paragraph, _EXPECTED["name"])


def test_contact_typography():
    identity = _synthetic_identity()
    doc = render(_synthetic_bank(), identity, _default_profile())
    paragraph = doc.paragraphs[1]
    assert identity.email in paragraph.text
    _assert_paragraph_matches(paragraph, _EXPECTED["contact"])


def test_contact_line_links_are_real_hyperlinks_not_url_text():
    """Regression coverage for a real spec 03 gap: the contact line used to
    print raw URLs as plain visible text. LinkedIn/GitHub/Portfolio/Scholar
    must be actual docx hyperlink relationships pointing at the
    identity.toml URLs, displayed as short labels, not the URL string."""
    identity = _synthetic_identity()
    doc = render(_synthetic_bank(), identity, _default_profile())
    paragraph = doc.paragraphs[1]

    hyperlinks = {hyperlink.text: hyperlink.url for hyperlink in paragraph.hyperlinks}
    assert hyperlinks == {
        "LinkedIn": identity.linkedin,
        "GitHub": identity.github,
        "Portfolio": identity.portfolio,
        "Scholar": identity.scholar,
    }


def test_contact_line_phone_and_email_are_plain_text_not_links():
    identity = _synthetic_identity()
    doc = render(_synthetic_bank(), identity, _default_profile())
    paragraph = doc.paragraphs[1]

    assert len(paragraph.hyperlinks) == 4
    assert identity.phone in paragraph.text
    assert identity.email in paragraph.text
    non_hyperlink_text = "".join(run.text for run in paragraph.runs)
    assert identity.phone in non_hyperlink_text
    assert identity.email in non_hyperlink_text


def test_contact_hyperlink_runs_are_blue_underlined_arial_12pt():
    identity = _synthetic_identity()
    doc = render(_synthetic_bank(), identity, _default_profile())
    paragraph = doc.paragraphs[1]

    assert paragraph.hyperlinks
    for hyperlink in paragraph.hyperlinks:
        for run in hyperlink.runs:
            assert run.font.name == "Arial"
            assert run.font.size == Pt(12)
            assert run.font.underline is True
            assert run.font.color.rgb == RGBColor.from_string("467886")


def test_no_raw_url_text_anywhere_in_document():
    identity = _synthetic_identity()
    doc = render(_synthetic_bank(), identity, _default_profile())
    for paragraph in doc.paragraphs:
        assert "https://" not in paragraph.text


def test_status_typography():
    identity = _synthetic_identity()
    doc = render(_synthetic_bank(), identity, _default_profile())
    paragraph = doc.paragraphs[2]
    assert identity.work_authorization_statement in paragraph.text
    _assert_paragraph_matches(paragraph, _EXPECTED["status"])


def test_section_header_typography():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    for header_text in ("Education & Certificates", "Work History", "Publications"):
        paragraph = _find_paragraph(doc, header_text)
        _assert_paragraph_matches(paragraph, _EXPECTED["section_header"])


def test_job_title_and_date_share_one_paragraph_with_correct_typography():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    (title_paragraph,) = _paragraphs_after(doc, "Work History", 1)
    assert "Test Engineer" in title_paragraph.text
    assert "TestCo" in title_paragraph.text
    _assert_paragraph_matches(title_paragraph, _EXPECTED["job_title_and_date"])


def test_bullet_typography():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    title_paragraph, summary_paragraph, bullet_1, bullet_2 = _paragraphs_after(
        doc, "Work History", 4
    )
    _assert_paragraph_matches(summary_paragraph, _EXPECTED["bullet"])
    _assert_paragraph_matches(bullet_1, _EXPECTED["bullet"])
    _assert_paragraph_matches(bullet_2, _EXPECTED["bullet"])


def test_publication_author_run_is_bold_others_are_not():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    header = _find_paragraph(doc, "Publications")
    (pub_paragraph,) = _paragraphs_after(doc, "Publications", 1)
    assert "Person, T." in pub_paragraph.text
    author_runs = [r for r in pub_paragraph.runs if "Person, T." in r.text]
    other_runs = [r for r in pub_paragraph.runs if "Person, T." not in r.text and r.text.strip()]
    assert author_runs
    for run in author_runs:
        assert run.font.bold is True
        assert run.font.name == "Arial"
        assert run.font.size == Pt(10.5)
    for run in other_runs:
        assert run.font.bold is False
    _assert_paragraph_matches(header, _EXPECTED["section_header"])


def test_page_margins_and_size():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    section = doc.sections[0]
    assert section.top_margin == _EXPECTED_MARGIN
    assert section.bottom_margin == _EXPECTED_MARGIN
    assert section.left_margin == _EXPECTED_MARGIN
    assert section.right_margin == _EXPECTED_MARGIN
    assert section.page_width == Twips(12240)
    assert section.page_height == Twips(15840)


def test_single_right_tab_stop_at_7_5_inches():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    (title_paragraph,) = _paragraphs_after(doc, "Work History", 1)
    tab_stops = list(title_paragraph.paragraph_format.tab_stops)
    assert len(tab_stops) == 1
    assert tab_stops[0].position == _EXPECTED_TAB_POSITION
    assert tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT


def test_no_paragraph_is_justified():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    for paragraph in doc.paragraphs:
        assert paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY


def test_education_date_uses_tab_stop_not_hardcoded_spaces():
    """Regression test: an earlier version of _render_education built the
    degree+status line with hardcoded spaces instead of a tab stop, so long
    degree text pushed the status text off alignment instead of staying
    pinned to the right margin. Uses a deliberately long degree/field/
    institution string, long enough that hardcoded spaces would visibly
    fail, to make sure this can't silently regress."""
    long_education = [
        Education(
            id="edu_long",
            degree="Master of Science",
            field=(
                "Highly Specialized Interdisciplinary Computational "
                "Engineering and Applied Mathematics"
            ),
            institution=(
                "A Very Long University Name That Goes On For Quite A "
                "While Indeed State University"
            ),
            gpa="4.0/4.0",
            status="May 2025",
            requires_degree_profiles=["ai_ml_engineer"],
        )
    ]
    doc = render(
        _synthetic_bank(education=long_education),
        _synthetic_identity(),
        _default_profile(),
    )
    (edu_paragraph,) = _paragraphs_after(doc, "Education & Certificates", 1)

    assert any("\t" in run.text for run in edu_paragraph.runs)
    assert "        " not in edu_paragraph.text  # no hardcoded-space fallback

    tab_stops = list(edu_paragraph.paragraph_format.tab_stops)
    assert len(tab_stops) == 1
    assert tab_stops[0].position == _EXPECTED_TAB_POSITION
    assert tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT
    _assert_paragraph_matches(edu_paragraph, _EXPECTED["bullet"])


# ---------------------------------------------------------------------------
# Structural rules from specs/03-renderer.md.
# ---------------------------------------------------------------------------


def test_education_section_order_default_is_first():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts.index("Education & Certificates") < texts.index("Work History")


def test_education_moves_to_bottom_when_profile_orders_it_last():
    profile = _default_profile(
        section_order=["work_history", "projects", "publications", "education"]
    )
    doc = render(_synthetic_bank(), _synthetic_identity(), profile)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts.index("Education & Certificates") > texts.index("Work History")
    assert texts.index("Education & Certificates") > texts.index("Publications")


def test_summary_section_omitted_by_default():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    texts = [p.text.strip() for p in doc.paragraphs]
    assert "Test Person" not in texts[3]  # blank line, not a summary paragraph
    for text in texts:
        assert text != "SUMMARY_MARKER_SHOULD_NOT_APPEAR"


def test_summary_section_included_when_profile_triggers_it():
    profile = _default_profile(
        include_summary=True,
        summary_text="Relocating and seeking a new industry",
    )
    doc = render(_synthetic_bank(), _synthetic_identity(), profile)
    texts = [p.text.strip() for p in doc.paragraphs]
    assert "Relocating and seeking a new industry" in texts


def test_role_summary_bullet_is_always_first_bullet_under_role():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    title_paragraph, summary_paragraph, bullet_1, bullet_2 = _paragraphs_after(
        doc, "Work History", 4
    )
    assert "Built the systems behind a test product" in summary_paragraph.text
    assert "Built a thing using a way that produced a result" in bullet_1.text


def test_current_role_renders_to_present():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    (title_paragraph,) = _paragraphs_after(doc, "Work History", 1)
    assert "Mar 2026" in title_paragraph.text
    assert "Present" in title_paragraph.text


def test_project_role_renders_no_date_line():
    doc = render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    project_header = _find_paragraph(doc, "Projects")
    (title_paragraph,) = _paragraphs_after(doc, "Projects", 1)
    assert "Test Side Project" in title_paragraph.text
    assert "2026" not in title_paragraph.text
    assert "Present" not in title_paragraph.text
    _assert_paragraph_matches(project_header, _EXPECTED["section_header"])


def test_rendering_does_not_mutate_template_file():
    before = hashlib.sha256(_TEMPLATE_PATH.read_bytes()).hexdigest()
    render(_synthetic_bank(), _synthetic_identity(), _default_profile())
    after = hashlib.sha256(_TEMPLATE_PATH.read_bytes()).hexdigest()
    assert before == after


# ---------------------------------------------------------------------------
# The actual golden test: full real bank, no tailoring, real identity.toml.
# ---------------------------------------------------------------------------


def test_full_real_bank_matches_golden_typography():
    from jobengine.resume.render import load_identity

    bank = load_bank()
    identity = load_identity(_IDENTITY_PATH)
    profile = _default_profile()
    doc = render(bank, identity, profile)

    _assert_paragraph_matches(_find_paragraph(doc, identity.full_name), _EXPECTED["name"])
    _assert_paragraph_matches(doc.paragraphs[1], _EXPECTED["contact"])
    _assert_paragraph_matches(doc.paragraphs[2], _EXPECTED["status"])

    contact_hyperlinks = {h.text: h.url for h in doc.paragraphs[1].hyperlinks}
    assert contact_hyperlinks == {
        "LinkedIn": identity.linkedin,
        "GitHub": identity.github,
        "Portfolio": identity.portfolio,
        "Scholar": identity.scholar,
    }
    for paragraph in doc.paragraphs:
        assert "https://" not in paragraph.text
    for header_text in (
        "Education & Certificates",
        "Work History",
        "Projects",
        "Publications",
    ):
        _assert_paragraph_matches(_find_paragraph(doc, header_text), _EXPECTED["section_header"])

    section = doc.sections[0]
    assert section.top_margin == _EXPECTED_MARGIN
    assert section.bottom_margin == _EXPECTED_MARGIN
    assert section.left_margin == _EXPECTED_MARGIN
    assert section.right_margin == _EXPECTED_MARGIN

    (title_paragraph,) = _paragraphs_after(doc, "Work History", 1)
    tab_stops = list(title_paragraph.paragraph_format.tab_stops)
    assert len(tab_stops) == 1
    assert tab_stops[0].position == _EXPECTED_TAB_POSITION
    assert tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT

    for paragraph in doc.paragraphs:
        assert paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY
