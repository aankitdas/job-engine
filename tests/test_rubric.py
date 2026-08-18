"""Tests for jobengine.rubric. See specs/08-rubric.md.

rubric.rules/measure/score do not exist yet; every test here is expected to
fail on collection until they're implemented. One failing fixture per hard
rule (R001-R013), per spec 08's definition of done, plus select_for_profile,
score.py's weighted formula, and one real-bank integration test.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from jobengine.resume.bank import (
    Bank,
    Bullet,
    Meta,
    Role,
    SummaryBullet,
    load_bank,
)
from jobengine.resume.pdf import render_pdf
from jobengine.resume.render import RenderProfile, load_identity, render
from jobengine.rubric import measure, rules, score

_IDENTITY_PATH = Path("identity.toml")


def _summary(
    id_: str = "s1",
    text: str = "Built systems.",
    keywords=None,
    status: str = "verified",
) -> SummaryBullet:
    return SummaryBullet(id=id_, text=text, keywords=keywords or [], status=status)


def _bullet(
    id_: str,
    text: str = "Built a thing that worked well.",
    keywords=None,
    profiles=None,
    status: str = "verified",
) -> Bullet:
    return Bullet(
        id=id_,
        status=status,
        what="a thing",
        how="carefully",
        result="it worked",
        text=text,
        keywords=keywords or [],
        profiles=profiles or ["ai_ml_engineer"],
    )


def _role(
    id_: str,
    *,
    kind: str = "full_time",
    start: str | None = "2022-01",
    end: str | None = "2023-01",
    bullets=None,
    summary=None,
) -> Role:
    return Role(
        id=id_,
        company="Acme" if kind != "project" else None,
        location="Remote" if kind != "project" else None,
        start=start,
        end=end,
        kind=kind,
        title={"default": "Engineer"},
        summary=summary or _summary(f"{id_}_s"),
        bullets=bullets
        if bullets is not None
        else [_bullet(f"{id_}_b1"), _bullet(f"{id_}_b2")],
    )


def _bank(roles: list[Role]) -> Bank:
    return Bank(meta=Meta(owner="test", updated="2026-01-01"), roles=roles)


# ---------------------------------------------------------------------------
# select_for_profile
# ---------------------------------------------------------------------------


def test_select_for_profile_keeps_only_tagged_bullets():
    role = _role(
        "r1",
        bullets=[
            _bullet("b1", profiles=["ai_ml_engineer"]),
            _bullet("b2", profiles=["software_engineer"]),
            _bullet("b3", profiles=["ai_ml_engineer", "software_engineer"]),
        ],
    )
    bank = _bank([role])
    selected = measure.select_for_profile(bank, "ai_ml_engineer")
    assert len(selected.roles) == 1
    ids = {b.id for b in selected.roles[0].bullets}
    assert ids == {"b1", "b3"}


def test_select_for_profile_drops_roles_with_zero_tagged_bullets():
    role_a = _role("r1", bullets=[_bullet("b1", profiles=["ai_ml_engineer"])])
    role_b = _role("r2", bullets=[_bullet("b2", profiles=["software_engineer"])])
    bank = _bank([role_a, role_b])
    selected = measure.select_for_profile(bank, "ai_ml_engineer")
    assert [r.id for r in selected.roles] == ["r1"]


def test_select_for_profile_original_bank_is_not_mutated():
    role = _role("r1", bullets=[_bullet("b1", profiles=["ai_ml_engineer"])])
    bank = _bank([role])
    measure.select_for_profile(bank, "software_engineer")
    assert len(bank.roles) == 1
    assert len(bank.roles[0].bullets) == 1


# ---------------------------------------------------------------------------
# R001: coverage
# ---------------------------------------------------------------------------


def test_r001_coverage_computed_as_set_intersection_stem_normalized():
    role = _role("r1", bullets=[_bullet("b1", keywords=["Kubernetes", "Airflow"])])
    bank = _bank([role])
    cov = measure.coverage(bank, ["Kubernetes", "Airflows", "Docker"])
    assert cov == pytest.approx(2 / 3)


def test_r001_fails_below_threshold():
    role = _role("r1", bullets=[_bullet("b1", keywords=["Kubernetes"])])
    bank = _bank([role])
    result = rules.check_r001(
        bank, required_keywords=["Kubernetes", "Airflow", "A/B testing"]
    )
    assert result is not None
    assert result.rule == "R001"


def test_r001_passes_at_or_above_threshold():
    role = _role("r1", bullets=[_bullet("b1", keywords=["Kubernetes", "Airflow"])])
    bank = _bank([role])
    result = rules.check_r001(
        bank, required_keywords=["Kubernetes", "Airflow", "A/B testing", "Docker"]
    )
    # 2/4 = 0.5 < 0.70, still a failure; use a passing ratio explicitly
    assert result is not None
    result_pass = rules.check_r001(bank, required_keywords=["Kubernetes", "Airflow"])
    assert result_pass is None


# ---------------------------------------------------------------------------
# R002: front-loading (real PDF geometry)
# ---------------------------------------------------------------------------


@pytest.fixture(scope="module")
def real_sample_pdf(tmp_path_factory) -> Path:
    bank = load_bank()
    identity = load_identity(_IDENTITY_PATH)
    profile = RenderProfile(
        section_order=["work_history", "projects", "education", "publications"],
        include_summary=False,
    )
    doc = render(bank, identity, profile)
    out_dir = tmp_path_factory.mktemp("rubric_pdf")
    docx_path = out_dir / "candidate.docx"
    doc.save(docx_path)
    return render_pdf(docx_path, out_dir)


def test_r002_front_load_measures_real_pdf_geometry(real_sample_pdf):
    ratio = measure.front_load(real_sample_pdf, ["Python", "LLM"])
    assert 0.0 <= ratio <= 1.0


def test_pdf_parsing_is_cached_per_file_hash_not_reparsed_every_call(
    real_sample_pdf, monkeypatch
):
    """Spec 08: "Cache the extraction per rendered file hash so repeated
    scoring is free." front_load/front_load_detail/line_count_from_pdf/
    page1_height must all share one parse of a given PDF, not each open
    and re-parse it independently."""
    import pdfplumber

    measure._PAGE_CACHE.clear()
    real_open = pdfplumber.open
    open_calls = []

    def _counting_open(*args, **kwargs):
        open_calls.append(args)
        return real_open(*args, **kwargs)

    monkeypatch.setattr(pdfplumber, "open", _counting_open)

    bank = load_bank()
    real_summary_text = bank.roles[0].summary.text

    measure.page1_height(real_sample_pdf)
    measure.front_load(real_sample_pdf, ["Python", "LLM"])
    measure.line_count_from_pdf(real_sample_pdf, real_summary_text)

    assert len(open_calls) == 1


def test_pdf_parsing_cache_is_keyed_by_content_not_path(
    real_sample_pdf, tmp_path, monkeypatch
):
    """Two different paths with byte-identical PDF content share one cache
    entry, matching spec 08's "per rendered file hash" wording rather than
    "per path" (relevant once job_resume_variants dedups identical
    selections onto one rendered file per spec 08's Storage section)."""
    import pdfplumber

    measure._PAGE_CACHE.clear()
    duplicate_path = tmp_path / "duplicate.pdf"
    duplicate_path.write_bytes(real_sample_pdf.read_bytes())

    real_open = pdfplumber.open
    open_calls = []

    def _counting_open(*args, **kwargs):
        open_calls.append(args)
        return real_open(*args, **kwargs)

    monkeypatch.setattr(pdfplumber, "open", _counting_open)

    measure.page1_height(real_sample_pdf)
    measure.page1_height(duplicate_path)

    assert len(open_calls) == 1


def test_r002_fails_below_threshold():
    result = rules.check_r002(front_load_ratio=0.5)
    assert result is not None
    assert result.rule == "R002"


def test_r002_passes_at_or_above_threshold():
    assert rules.check_r002(front_load_ratio=0.75) is None
    assert rules.check_r002(front_load_ratio=1.0) is None


# ---------------------------------------------------------------------------
# R003: 3 to 8 bullets per role including summary
# ---------------------------------------------------------------------------


def test_r003_fails_role_with_too_few_bullets():
    role = _role("r1", bullets=[])  # 1 (summary) total
    bank = _bank([role])
    result = rules.check_r003(bank)
    assert result is not None
    assert result.rule == "R003"


def test_r003_fails_role_with_too_many_bullets():
    role = _role("r1", bullets=[_bullet(f"b{i}") for i in range(8)])  # 9 total
    bank = _bank([role])
    result = rules.check_r003(bank)
    assert result is not None


def test_r003_passes_role_in_range():
    role = _role("r1", bullets=[_bullet("b1"), _bullet("b2")])  # 3 total
    bank = _bank([role])
    assert rules.check_r003(bank) is None


# ---------------------------------------------------------------------------
# R004: exactly one summary bullet (structurally guaranteed by the schema)
# ---------------------------------------------------------------------------


def test_r004_passes_by_construction():
    role = _role("r1")
    bank = _bank([role])
    assert rules.check_r004(bank) is None


# ---------------------------------------------------------------------------
# R005: at most one period per bullet
# ---------------------------------------------------------------------------


def test_r005_fails_bullet_with_two_periods():
    role = _role("r1", bullets=[_bullet("b1", text="Did a thing. Then did another.")])
    bank = _bank([role])
    result = rules.check_r005(bank)
    assert result is not None
    assert result.rule == "R005"


def test_r005_passes_bullet_with_one_period():
    role = _role("r1", bullets=[_bullet("b1", text="Did a thing.")])
    bank = _bank([role])
    assert rules.check_r005(bank) is None


# ---------------------------------------------------------------------------
# R006: renders in 3 lines or fewer (real PDF geometry)
# ---------------------------------------------------------------------------


def test_r006_measures_real_bullet_line_count(real_sample_pdf):
    bank = load_bank()
    long_bullet = None
    for role in bank.roles:
        for b in role.bullets:
            if long_bullet is None or len(b.text) > len(long_bullet.text):
                long_bullet = b
    assert long_bullet is not None
    line_count = measure.line_count_from_pdf(real_sample_pdf, long_bullet.text)
    assert line_count >= 1


def test_r006_fails_bullet_rendering_over_three_lines():
    result = rules.check_r006(line_counts={"b1": 4})
    assert result is not None
    assert result.rule == "R006"


def test_r006_passes_bullet_at_three_lines_or_fewer():
    assert rules.check_r006(line_counts={"b1": 3, "b2": 1}) is None


# ---------------------------------------------------------------------------
# R007: past tense
# ---------------------------------------------------------------------------


def test_r007_fails_present_tense_bullet():
    role = _role("r1", bullets=[_bullet("b1", text="Manages a team of five.")])
    bank = _bank([role])
    result = rules.check_r007(bank)
    assert result is not None
    assert result.rule == "R007"


def test_r007_passes_past_tense_bullet():
    role = _role("r1", bullets=[_bullet("b1", text="Managed a team of five.")])
    bank = _bank([role])
    assert rules.check_r007(bank) is None


# ---------------------------------------------------------------------------
# R008: no first-person pronouns
# ---------------------------------------------------------------------------


def test_r008_fails_bullet_with_first_person_pronoun():
    role = _role("r1", bullets=[_bullet("b1", text="I built a thing that worked.")])
    bank = _bank([role])
    result = rules.check_r008(bank)
    assert result is not None
    assert result.rule == "R008"


def test_r008_does_not_false_positive_on_substring():
    # "Iterated"/"Improved" must not trigger on the "I" word-boundary check.
    role = _role("r1", bullets=[_bullet("b1", text="Iterated on the design weekly.")])
    bank = _bank([role])
    assert rules.check_r008(bank) is None


def test_r008_passes_clean_bullet():
    role = _role("r1", bullets=[_bullet("b1", text="Built a thing that worked well.")])
    bank = _bank([role])
    assert rules.check_r008(bank) is None


# ---------------------------------------------------------------------------
# R009: reverse chronological order
# ---------------------------------------------------------------------------


def test_r009_fails_out_of_order_roles():
    role_old = _role("r_old", start="2018-01", end="2019-01")
    role_new = _role("r_new", start="2022-01", end="2023-01")
    bank = _bank([role_old, role_new])  # older role listed first: wrong order
    result = rules.check_r009(bank)
    assert result is not None
    assert result.rule == "R009"


def test_r009_passes_correctly_ordered_roles():
    role_new = _role("r_new", start="2022-01", end="2023-01")
    role_old = _role("r_old", start="2018-01", end="2019-01")
    bank = _bank([role_new, role_old])
    assert rules.check_r009(bank) is None


def test_r009_ignores_project_roles_with_no_dates():
    role_new = _role("r_new", start="2022-01", end="2023-01")
    project = _role("r_proj", kind="project", start=None, end=None)
    bank = _bank([role_new, project])
    assert rules.check_r009(bank) is None


def test_r009_tolerates_genuinely_overlapping_roles_in_either_order():
    # role_sei/role_unl-style real data: 2021-10 to 2023-08 and 2021-05 to
    # 2023-06 genuinely overlap despite different start dates. Neither
    # order is a violation: R009 only flags a role appearing before an
    # earlier, non-overlapping role, not any difference in start date.
    role_a = _role("r_a", start="2021-10", end="2023-08")
    role_b = _role("r_b", start="2021-05", end="2023-06")
    assert rules.check_r009(_bank([role_a, role_b])) is None
    assert rules.check_r009(_bank([role_b, role_a])) is None


def test_r009_still_fails_non_overlapping_roles_in_wrong_order():
    role_old = _role("r_old", start="2018-01", end="2019-01")
    role_new = _role("r_new", start="2022-01", end="2023-01")
    assert rules.check_r009(_bank([role_old, role_new])) is not None
    assert rules.check_r009(_bank([role_new, role_old])) is None


# ---------------------------------------------------------------------------
# R010: typography matches the golden spec
# ---------------------------------------------------------------------------


def test_r010_passes_a_real_render(tmp_path):
    bank = load_bank()
    identity = load_identity(_IDENTITY_PATH)
    profile = RenderProfile(section_order=["work_history"], include_summary=False)
    doc = render(bank, identity, profile)
    docx_path = tmp_path / "candidate.docx"
    doc.save(docx_path)
    violations = measure.measure_typography(docx_path)
    assert violations == []


def test_r010_fails_wrong_margins(tmp_path):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.sections[0].left_margin = Inches(1.0)
    docx_path = tmp_path / "bad.docx"
    doc.save(docx_path)
    violations = measure.measure_typography(docx_path)
    assert violations != []


# ---------------------------------------------------------------------------
# R011: single column
# ---------------------------------------------------------------------------


def test_r011_passes_a_real_render(tmp_path):
    bank = load_bank()
    identity = load_identity(_IDENTITY_PATH)
    profile = RenderProfile(section_order=["work_history"], include_summary=False)
    doc = render(bank, identity, profile)
    docx_path = tmp_path / "candidate.docx"
    doc.save(docx_path)
    assert measure.is_single_column(docx_path) is True


def test_r011_fails_docx_with_a_table(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_table(rows=1, cols=2)
    docx_path = tmp_path / "table.docx"
    doc.save(docx_path)
    assert measure.is_single_column(docx_path) is False


# ---------------------------------------------------------------------------
# R012: zero speculative bullets
# ---------------------------------------------------------------------------


def test_r012_fails_speculative_bullet():
    role = _role("r1", bullets=[_bullet("b1", status="speculative")])
    bank = _bank([role])
    result = rules.check_r012(bank)
    assert result is not None
    assert result.rule == "R012"


def test_r012_fails_speculative_summary():
    role = _role("r1", summary=_summary(status="speculative"))
    bank = _bank([role])
    result = rules.check_r012(bank)
    assert result is not None


def test_r012_passes_all_verified():
    role = _role("r1", bullets=[_bullet("b1", status="verified")])
    bank = _bank([role])
    assert rules.check_r012(bank) is None


# ---------------------------------------------------------------------------
# R013: slop linter passes with zero errors
# ---------------------------------------------------------------------------


def test_r013_fails_when_slop_linter_finds_an_error():
    # A first-person pronoun trips slop_lint's own H007, an error not a
    # warning. Two bullets so the role's own 3-8 total (slop_lint's H004)
    # doesn't independently fail and confound the assertion.
    role = _role(
        "r1",
        bullets=[
            _bullet("b1", text="I built a thing that worked."),
            _bullet("b2", text="Shipped a second thing too."),
        ],
    )
    bank = _bank([role])
    result = rules.check_r013(bank)
    assert result is not None
    assert result.rule == "R013"


def test_r013_passes_clean_bank():
    role = _role(
        "r1",
        bullets=[
            _bullet("b1", text="Built a thing that worked well."),
            _bullet("b2", text="Shipped a second thing too."),
        ],
    )
    bank = _bank([role])
    assert rules.check_r013(bank) is None


# ---------------------------------------------------------------------------
# score.py
# ---------------------------------------------------------------------------


def test_score_is_weighted_sum_of_five_components(monkeypatch):
    monkeypatch.setattr(score, "_coverage_component", lambda *a, **k: 1.0)
    monkeypatch.setattr(score, "_front_load_component", lambda *a, **k: 1.0)
    monkeypatch.setattr(score, "_density_component", lambda *a, **k: 1.0)
    monkeypatch.setattr(score, "_multi_keyword_component", lambda *a, **k: 1.0)
    monkeypatch.setattr(score, "_page_penalty_component", lambda *a, **k: 0.0)
    result = score.compute_score(
        bank=_bank([_role("r1")]),
        required_keywords=["x"],
        preferred_keywords=[],
        pdf_path=Path("unused.pdf"),
        page_count=1,
    )
    assert result == pytest.approx(100.0)


def test_density_component_is_distinct_keyword_hits_over_first_role_word_count():
    role = _role(
        "r1",
        bullets=[
            _bullet("b1", text="Built Kubernetes pipelines for the whole team quickly")
        ],
    )
    density = score._density_component(role, ["Kubernetes", "Airflow"])
    # 1 distinct hit ("Kubernetes") over the first role's word count.
    word_count = len(role.summary.text.split()) + len(
        ["Built", "Kubernetes", "pipelines", "for", "the", "whole", "team", "quickly"]
    )
    assert density == pytest.approx(1 / word_count)


def test_multi_keyword_component_counts_bullets_across_whole_resume():
    role1 = _role(
        "r1",
        bullets=[
            _bullet("b1", keywords=["Kubernetes", "Airflow"]),
            _bullet("b2", keywords=[]),
        ],
    )
    role2 = _role(
        "r2",
        bullets=[
            _bullet("b3", keywords=["Kubernetes"]),
            _bullet("b4", keywords=["Docker"]),
        ],
    )
    bank = _bank([role1, role2])
    ratio = score._multi_keyword_component(bank, ["Kubernetes", "Airflow", "Docker"])
    # Only b1 carries >= 2 target keywords, out of 4 total bullets.
    assert ratio == pytest.approx(1 / 4)


# ---------------------------------------------------------------------------
# Integration: real bank, real render, real PDF, real (non-persisted) C3
# extraction output against a fixed keyword list standing in for a live job.
# ---------------------------------------------------------------------------


def test_score_resume_runs_end_to_end_against_the_real_bank(real_sample_pdf):
    bank = load_bank()
    candidate = measure.select_for_profile(bank, "ai_ml_engineer")
    result = rules.score_resume(
        bank=candidate,
        profile="ai_ml_engineer",
        docx_path=real_sample_pdf.with_suffix(".docx"),
        pdf_path=real_sample_pdf,
        required_keywords=["Python", "LLM", "RAG", "Kubernetes", "Airflow"],
        preferred_keywords=[],
    )
    assert 0.0 <= result.score <= 100.0
    assert isinstance(result.passed, bool)
    assert "coverage" in result.measurements


# --- has_unrecoverable_rubric_failure: spec 08's P4 soft/hard split
# (D42) -- R001-only is a "soft" deficit eligible for human override at
# approve() time; any other hard rule is a genuine document defect and
# is never overridable.


def test_no_failures_is_not_unrecoverable():
    assert rules.has_unrecoverable_rubric_failure([]) is False


def test_r001_only_is_not_unrecoverable():
    assert rules.has_unrecoverable_rubric_failure(["R001"]) is False


def test_r001_plus_another_rule_is_unrecoverable():
    assert rules.has_unrecoverable_rubric_failure(["R001", "R003"]) is True


def test_r006_alone_is_unrecoverable():
    assert rules.has_unrecoverable_rubric_failure(["R006"]) is True


def test_multiple_non_r001_rules_are_unrecoverable():
    assert rules.has_unrecoverable_rubric_failure(["R010", "R013"]) is True


# --- RULE_INFO: human-readable name/description per rule, sourced from
# specs/08-rubric.md (D46), the one place the review page's rule
# explanations come from so the UI and the spec can't drift apart.
# R012/R013 are deliberately not yet in RULE_INFO -- spec 08's own text
# for both is too terse to render usefully (D46), pending real text from
# the user rather than an invented paraphrase.


def test_rule_info_has_eleven_confirmed_entries():
    assert set(rules.RULE_INFO.keys()) == {
        "R001",
        "R002",
        "R003",
        "R004",
        "R005",
        "R006",
        "R007",
        "R008",
        "R009",
        "R010",
        "R011",
    }


def test_every_rule_info_entry_has_a_nonempty_name_and_description():
    for rule_id, info in rules.RULE_INFO.items():
        assert info.name.strip(), rule_id
        assert info.description.strip(), rule_id


def test_r001_description_states_the_070_threshold():
    assert "0.70" in rules.RULE_INFO["R001"].description


def test_r010_description_names_arial():
    assert "Arial" in rules.RULE_INFO["R010"].description
