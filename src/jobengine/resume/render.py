"""Docx renderer. See specs/03-renderer.md.

Typography is applied directly, per run and paragraph, from _TYPOGRAPHY
below. Nothing is left to inherit from a Word style default; the template's
own styles.xml has a stale-default bug (one bullet at 10pt instead of
10.5pt) that is exactly what happens when something isn't set explicitly.

Every document is built fresh via docx.Document(); resume/templates/
golden.docx is never opened for writing, so rule 6 (never mutate the
template) holds by construction, not by care.
"""

from __future__ import annotations

import tomllib
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

from jobengine.resume.bank import Bank, Education, Publication, Role

DEFAULT_IDENTITY_PATH = Path("identity.toml")

_FONT = "Arial"
_MARGIN = Twips(720)
_TAB_POSITION = Twips(10800)
# Matches the template's own Hyperlink character style (w:color w:val="467886").
_HYPERLINK_COLOR = "467886"


@dataclass(frozen=True)
class _Style:
    size: Pt
    bold: bool = False
    italic: bool = False
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT
    line_spacing: float = 1.5


_TYPOGRAPHY = {
    "name": _Style(
        Pt(14), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15
    ),
    "contact": _Style(Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15),
    "status": _Style(Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15),
    "section_header": _Style(Pt(10.5), bold=True, line_spacing=1.5),
    "job_title_and_date": _Style(Pt(10.5), italic=True, line_spacing=1.5),
    "bullet": _Style(Pt(10.5), line_spacing=1.5),
    "publication": _Style(Pt(10.5), line_spacing=1.5),
}


@dataclass
class Identity:
    full_name: str
    email: str
    phone: str
    city: str
    state: str
    linkedin: str
    github: str
    portfolio: str
    scholar: str
    work_authorization_statement: str


def load_identity(path: Path = DEFAULT_IDENTITY_PATH) -> Identity:
    raw = tomllib.loads(path.read_text())
    contact = raw["contact"]
    work_authorization = raw["work_authorization"]
    return Identity(
        full_name=contact["full_name"],
        email=contact["email"],
        phone=contact["phone"],
        city=contact["city"],
        state=contact["state"],
        linkedin=contact["linkedin"],
        github=contact["github"],
        portfolio=contact["portfolio"],
        scholar=contact["scholar"],
        work_authorization_statement=work_authorization["statement"],
    )


@dataclass
class RenderProfile:
    """Stand-in for E1's not-yet-built profile registry. render() only
    needs section order and whether the summary section is triggered; where
    that data ultimately comes from is E1's decision, not render.py's."""

    section_order: list[str]
    include_summary: bool = False
    summary_text: str | None = None


def _set_run_style(run, style_key: str) -> None:
    style = _TYPOGRAPHY[style_key]
    run.font.name = _FONT
    run.font.size = style.size
    run.font.bold = style.bold
    run.font.italic = style.italic


def _add_hyperlink_run(paragraph, url: str, text: str, style_key: str) -> None:
    """A real docx hyperlink relationship, not visible URL text. python-docx
    has no high-level write API for this: build the run normally (so it
    goes through the same _set_run_style as everything else), then move its
    <w:r> element inside a <w:hyperlink r:id="..."> that points at a
    relationship added via part.relate_to()."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    run = paragraph.add_run(text)
    _set_run_style(run, style_key)
    run.font.color.rgb = RGBColor.from_string(_HYPERLINK_COLOR)
    run.font.underline = True

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_element = run._r
    run_element.getparent().remove(run_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _add_paragraph(doc: Document, text: str, style_key: str):
    style = _TYPOGRAPHY[style_key]
    paragraph = doc.add_paragraph()
    paragraph.alignment = style.alignment
    paragraph.paragraph_format.line_spacing = style.line_spacing
    if text:
        run = paragraph.add_run(text)
        _set_run_style(run, style_key)
    return paragraph


def _add_blank(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = _TYPOGRAPHY["bullet"].line_spacing


def _add_section_header(doc: Document, text: str) -> None:
    _add_paragraph(doc, text, "section_header")


_MONTH_ABBREVIATIONS = {
    "01": "Jan",
    "02": "Feb",
    "03": "Mar",
    "04": "Apr",
    "05": "May",
    "06": "Jun",
    "07": "Jul",
    "08": "Aug",
    "09": "Sep",
    "10": "Oct",
    "11": "Nov",
    "12": "Dec",
}


def _format_month_year(yyyy_mm: str) -> str:
    year, month = yyyy_mm.split("-")
    return f"{_MONTH_ABBREVIATIONS[month]} {year}"


def _date_range_text(role: Role) -> str | None:
    if role.start is None:
        return None
    start = _format_month_year(role.start)
    end = _format_month_year(role.end) if role.end else "Present"
    return f"{start} to {end}"


def _add_right_tab_paragraph(
    doc: Document, before: str, after: str, style_key: str
) -> None:
    """A paragraph whose second half is pinned to the single 7.5in right tab
    stop, not hardcoded spaces. Hardcoded spaces only line up for one exact
    text length; a tab stop holds the right margin regardless of how long
    `before` runs, which is the whole point of a tab stop over spaces."""
    style = _TYPOGRAPHY[style_key]
    paragraph = doc.add_paragraph()
    paragraph.alignment = style.alignment
    paragraph.paragraph_format.line_spacing = style.line_spacing
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        _TAB_POSITION, WD_TAB_ALIGNMENT.RIGHT
    )
    run = paragraph.add_run(f"{before}\t{after}")
    _set_run_style(run, style_key)


def _add_job_title_paragraph(doc: Document, role: Role, title: str) -> None:
    title_text = (
        f"{title} at {role.company}, {role.location}" if role.company else title
    )
    date_text = _date_range_text(role)

    if date_text is not None:
        _add_right_tab_paragraph(doc, title_text, date_text, "job_title_and_date")
    else:
        _add_paragraph(doc, title_text, "job_title_and_date")


def _render_role(doc: Document, role: Role) -> None:
    title = role.title.get("default", "")
    _add_job_title_paragraph(doc, role, title)
    _add_paragraph(doc, f"• {role.summary.text}", "bullet")
    for bullet in role.bullets:
        _add_paragraph(doc, f"• {bullet.text}", "bullet")


def _render_role_group(doc: Document, header: str, roles: list[Role]) -> None:
    if not roles:
        return
    _add_section_header(doc, header)
    for index, role in enumerate(roles):
        _render_role(doc, role)
        if index < len(roles) - 1:
            _add_blank(doc)


def _render_education(doc: Document, education: list[Education]) -> None:
    if not education:
        return
    _add_section_header(doc, "Education & Certificates")
    for edu in education:
        parts = f"{edu.degree}, {edu.field}, {edu.institution}"
        if edu.gpa:
            parts += f" (GPA - {edu.gpa})"
        _add_right_tab_paragraph(doc, f"• {parts}", edu.status, "bullet")


def _render_publications(doc: Document, publications: list[Publication]) -> None:
    if not publications:
        return
    _add_section_header(doc, "Publications")
    style = _TYPOGRAPHY["publication"]
    for pub in publications:
        paragraph = doc.add_paragraph()
        paragraph.alignment = style.alignment
        paragraph.paragraph_format.line_spacing = style.line_spacing

        author_run = paragraph.add_run(f"• {pub.authors_bold} ")
        _set_run_style(author_run, "publication")
        author_run.font.bold = True

        rest_run = paragraph.add_run(f"{pub.text} {pub.venue}.")
        _set_run_style(rest_run, "publication")
        rest_run.font.bold = False


def _set_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = _MARGIN
    section.bottom_margin = _MARGIN
    section.left_margin = _MARGIN
    section.right_margin = _MARGIN


def _add_contact_line(doc: Document, identity: Identity) -> None:
    """Phone and email are plain black text, matching the template. The
    other four are short labels, each a real hyperlink relationship to the
    identity.toml URL, never the raw URL as visible text."""
    style = _TYPOGRAPHY["contact"]
    paragraph = doc.add_paragraph()
    paragraph.alignment = style.alignment
    paragraph.paragraph_format.line_spacing = style.line_spacing

    def plain(text: str) -> None:
        run = paragraph.add_run(text)
        _set_run_style(run, "contact")

    def separator() -> None:
        plain(" | ")

    plain(identity.phone)
    separator()
    plain(identity.email)
    separator()
    _add_hyperlink_run(paragraph, identity.linkedin, "LinkedIn", "contact")
    separator()
    _add_hyperlink_run(paragraph, identity.github, "GitHub", "contact")
    separator()
    _add_hyperlink_run(paragraph, identity.portfolio, "Portfolio", "contact")
    separator()
    _add_hyperlink_run(paragraph, identity.scholar, "Scholar", "contact")


def render(bank: Bank, identity: Identity, profile: RenderProfile) -> Document:
    doc = Document()
    _set_margins(doc)

    _add_paragraph(doc, identity.full_name, "name")
    _add_contact_line(doc, identity)
    status_text = f"{identity.work_authorization_statement} | {identity.state}"
    _add_paragraph(doc, status_text, "status")
    _add_blank(doc)

    if profile.include_summary and profile.summary_text:
        _add_paragraph(doc, profile.summary_text, "bullet")

    non_project_roles = [r for r in bank.roles if r.kind != "project"]
    project_roles = [r for r in bank.roles if r.kind == "project"]

    section_renderers = {
        "education": lambda: _render_education(doc, bank.education),
        "work_history": lambda: _render_role_group(
            doc, "Work History", non_project_roles
        ),
        "projects": lambda: _render_role_group(doc, "Projects", project_roles),
        "publications": lambda: _render_publications(doc, bank.publications),
    }
    for section in profile.section_order:
        section_renderers[section]()

    return doc
